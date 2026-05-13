"""
Convert PROJECT_SUMMARY.md to PROJECT_SUMMARY.docx
Run: python3 generate_word_doc.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.request
import io

MD_PATH = Path(__file__).parent.parent / "PROJECT_SUMMARY.md"
DOCX_PATH = Path(__file__).parent.parent / "PROJECT_SUMMARY.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "CCCCCC")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run_with_inline(paragraph, text: str, base_bold=False, base_italic=False):
    """Handles **bold**, `code`, and plain text in a mixed line."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|_[^_]+_)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("_") and part.endswith("_"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            r = paragraph.add_run(part)
        r.bold = base_bold or r.bold
        r.italic = base_italic or r.italic


# ── Main builder ─────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Drop separator row (---|---...)
        data_rows = [r for r in table_rows if not re.match(r"^\s*\|?[-| :]+\|?\s*$", r)]
        if not data_rows:
            table_rows = []
            in_table = False
            return

        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in re.split(r"(?<!\\)\|", row.strip().strip("|"))]
            parsed.append(cells)

        col_count = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=col_count)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = "Table Grid"

        for ri, row in enumerate(parsed):
            for ci, cell_text in enumerate(row):
                if ci >= col_count:
                    break
                cell = tbl.rows[ri].cells[ci]
                set_cell_border(cell)
                if ri == 0:
                    set_cell_bg(cell, "E8F0FE")
                p = cell.paragraphs[0]
                add_run_with_inline(p, cell_text, base_bold=(ri == 0))

        doc.add_paragraph()  # spacing after table
        table_rows = []
        in_table = False

    def flush_code(lang, lines_):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        code_text = "\n".join(lines_)
        run = p.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # Light grey shading on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        pPr.append(shd)

    skip_mermaid_image = False

    while i < len(lines):
        line = lines[i]

        # ── Code block open/close ──
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
                # Skip mermaid blocks (diagram already shown as image)
                if code_lang == "mermaid":
                    skip_mermaid_image = True
                else:
                    skip_mermaid_image = False
            else:
                in_code_block = False
                if not skip_mermaid_image:
                    flush_code(code_lang, code_lines)
                else:
                    # Add a note instead
                    p = doc.add_paragraph()
                    r = p.add_run("[ Mermaid architecture diagram — see image above ]")
                    r.italic = True
                    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table rows ──
        if line.strip().startswith("|"):
            if not in_table:
                in_table = True
            table_rows.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Blank line ──
        if not line.strip():
            i += 1
            continue

        # ── HR ──
        if re.match(r"^-{3,}$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "4A90D9")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Headings ──
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.runs[0].font.color.rgb = RGBColor(0x1A, 0x4F, 0x8A)
            elif level == 2:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.color.rgb = RGBColor(0x1A, 0x4F, 0x8A)
            elif level == 3:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            else:
                p = doc.add_heading(text, level=3)
            i += 1
            continue

        # ── Mermaid.ink image ──
        m = re.match(r"!\[.*?\]\((https://mermaid\.ink/img/[^\)]+)\)", line)
        if m:
            url = m.group(1)
            try:
                with urllib.request.urlopen(url, timeout=15) as resp:
                    img_data = resp.read()
                doc.add_picture(io.BytesIO(img_data), width=Inches(6))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p = doc.add_paragraph()
                r = p.add_run(f"[ Architecture diagram image: {url} ]")
                r.italic = True
                r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            i += 1
            continue

        # ── Generic image (skip) ──
        if re.match(r"!\[.*?\]\(.*?\)", line):
            i += 1
            continue

        # ── Blockquote ──
        if line.strip().startswith("> "):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_run_with_inline(p, text, base_italic=True)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "4A90D9")
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Unordered list ──
        m = re.match(r"^(\s*)[*\-]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3 + indent * 0.2)
            add_run_with_inline(p, text)
            i += 1
            continue

        # ── Ordered list ──
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            text = m.group(1)
            p = doc.add_paragraph(style="List Number")
            add_run_with_inline(p, text)
            i += 1
            continue

        # ── Normal paragraph ──
        p = doc.add_paragraph()
        add_run_with_inline(p, line.strip())
        i += 1

    if in_table:
        flush_table()

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    build_doc()
